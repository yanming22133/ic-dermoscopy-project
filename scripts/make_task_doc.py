"""生成 Task1/2/3 要求 + 评分标准 Word 文档"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.size = Pt(11)
style.font.name = 'Calibri'

def h1(text):
    doc.add_heading(text, level=1)
def h2(text):
    doc.add_heading(text, level=2)
def para(text):
    doc.add_paragraph(text)
def bullet(text):
    doc.add_paragraph(text, style='List Bullet')
def table(rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]), style='Light Grid Accent 1')
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            t.rows[i].cells[j].text = cell
    doc.add_paragraph()

doc.add_heading('IC DSI Summer School', level=0)
doc.add_heading('Project Requirements & Scoring Criteria', level=1)
doc.add_heading('项目要求与评分标准', level=1)
para('Project: Segmentation, Attribute Detection, and Anchored Findings Reports')
para('Group 2 | 2026-07-21 ~ 07-31')

# Task 1
h1('Task 1 -- Lesion Segmentation / 病灶分割')
h2('Input / 输入')
bullet('Dermoscopy image -- 640x480 RGB')
h2('Output / 输出')
bullet('Binary lesion mask -- {id}.jpg, L mode, 0/255')
h2('Evaluation Metrics / 评估指标')
bullet('Dice Coefficient (F1) -- overlap between prediction & GT')
bullet('IoU / Jaccard Index -- intersection over union')
bullet('95% Hausdorff Distance -- boundary distance (95th percentile)')
h2('Weight / 权重')
para('25%')
h2('Notes / 说明')
bullet('Accurate segmentation supports downstream Task3 evidence extraction')
bullet('Preprocessing: DullRazor hair removal + Shades of Gray + normalize + resize (tutorial p35)')
bullet('Augmentation: Rotation, Flip, Color Jitter, Scaling, Cropping, Elastic (tutorial p35)')
bullet('Loss: Dice Loss (seg) + BCE (cls) (tutorial p35)')
bullet('"Find the optimal backbone model" (tutorial p38)')

doc.add_page_break()

# Task 2
h1('Task 2 -- Attribute Detection / 属性检测')
h2('Output / 输出')
bullet('5 attribute masks (binary) -- {id}/{attr}.png, L, 0/255')
bullet('Per-attribute presence: present / absent / uncertain')
h2('Five Attributes / 五属性')
table([
    ['Attribute', 'Presence Rate'],
    ['Pigment Network', '~46%'],
    ['Negative Network', '~10% (sparse)'],
    ['Streaks', '~8% (sparse)'],
    ['Milia-like Cysts', '~20%'],
    ['Globules', '~22%'],
])
h2('Presence Definition / presence 定义 (tutorial p40)')
bullet('p_attr = mean(sigmoid(logits_attr)) over lesion ROI (recommended)')
bullet('Status: p>=0.60 present / p<=0.40 absent / else uncertain')
h2('Evaluation / 评估')
bullet('Per-attribute F1 / AUPRC (report separately -- best practice)')
bullet('Per-attribute mask Dice / IoU')
bullet('Focus on recall for sparse classes')
h2('Weight / 权重')
para('25%')
h2('Notes / 说明')
bullet('Loss = Loss_seg + 0.5*Loss_cls (tutorial p35)')
bullet('Sparse attribute masks (tutorial p53)')
bullet('Task 2: From Attribute Mask to Presence Label (tutorial p54)')

doc.add_page_break()

# Task 3
h1('Task 3 -- Anchored Findings Report / 锚定报告')
h2('Input / 输入')
bullet('Task1 mask -> lesion size + border irregularity')
bullet('Task2 presence + masks -> evidence description')
h2('Output / 输出')
bullet('JSON -- match example_result schema: image_id / split / model_version / attributes_order[5] / outputs.presence.{attr}.{prob,status}')
bullet('English report text -- all 5 terms, statuses match JSON')
bullet('Summary_reports_text.csv -- image_id, findings')
h2('Tutorial Thresholds / 教程阈值 (p40) -- hardcoded')
bullet('Attribute status: p>=0.60 present / p<=0.40 absent / else uncertain')
bullet('border_irregularity = perimeter^2 / (4pi*area); >=1.60 irregular; <1.60 regular')
bullet('lesion_area_ratio = lesion_pixels / total_pixels; <0.08 small / 0.08-0.25 moderate / >0.25 large')
h2('Report Text Format / 报告文本格式')
bullet('"The lesion is {size} with {border} borders. Pigment network is {status}; Negative network is {status}; Streaks are {status}; Milia-like cysts are {status}; Globules are {status}."')
bullet('Leading space in findings field')
bullet('is/are: "Pigment network is" / "Streaks are" / "Globules are"')
h2('Consistency Checklist / 一致性校验 (tutorial p40)')
bullet('1. All five required terms present')
bullet('2. Text status == JSON status')
bullet('3. Evidence aligns with predicted masks')
bullet('4. JSON schema compliance, valid numeric ranges, proper medical terminology')
h2('Weight / 权重')
para('20%')
h2('Constraints / 约束')
bullet('Must include all 5 medical terms')
bullet('Statuses must match JSON')
bullet('Evidence must align with masks')
bullet('No copying neighbours facts (tutorial p43)')

doc.add_page_break()

# Bonus
h1('Bonus -- CLIP Retrieval + RAG / CLIP检索+RAG')
h2('Retrieval / 检索 (tutorial p41-42)')
bullet('CLIP ViT-B/32, frozen, no fine-tuning')
bullet('Offline: FAISS cosine index over train images')
bullet('Online: Top-K neighbours per test image')
bullet('Crop lesion with Task1 mask before retrieval')
bullet('DINOv2 comparison optional')
h2('RAG (tutorial p43)')
bullet('Neighbour presence prior smoothing')
bullet('Grounded: no copying neighbour facts')
h2('Scoring / 评分 (tutorial p46)')
bullet('Audit Completeness: all images, all neighbour IDs saved')
bullet('Sanity: no duplicate IDs, consistent K, valid similarities')
bullet('Retrieval Impact: prove retrieval helps reports')
h2('Output / 输出')
bullet('test_bonus_clip.csv: query_image, neighbor_id, similarity')
bullet('neighbor_id format: data/images/{id}.jpg')
h2('Weight / 权重')
para('~30% (25+25+20=70, bonus unstated %)')

doc.add_page_break()

# Scoring Overview
h1('Scoring Breakdown / 评分总览')
table([
    ['Item', 'Weight', 'Key Metrics'],
    ['Task 1 Segmentation', '25%', 'Dice, IoU, 95% Hausdorff'],
    ['Task 2 Attribute', '25%', 'Per-attr F1/AUPRC, mask Dice/IoU'],
    ['Task 3 Report Consistency', '20%', '5 terms, status match, evidence align'],
    ['Bonus Retrieval', '~30%', 'Audit completeness, sanity, impact'],
])
para('Technical Report (<=12 pages) + Presentation (15+5 min) influence Best Overall & Best Segmentation awards.')

# Deliverables
h1('Deliverables / 交付物')
bullet('task1/ -- {id}.jpg, L, 0/255')
bullet('task2/ -- {id}/{attr}.png, L, 0/255 (attr singular: milia_like_cyst)')
bullet('task3/json/ -- {id}.json; task3/Summary_reports_text.csv')
bullet('bonus/ -- test_bonus_clip.csv')
bullet('code/ -- reproducible, fixed seed, run_inference.py')
bullet('Technical Report -- <=12 pages, Lit review/Methods/Results/Future')
bullet('Each group submits ONE set / 每组一套')

# Key Dates
h1('Key Dates / 关键日期')
table([
    ['Date', 'Event'],
    ['7/27', 'Submit Literature Review'],
    ['7/30 13:00', 'Test Set Released'],
    ['7/30 19:00', 'Submit Code + Results + Report (6h window)'],
    ['7/31', 'Final Presentation + Awards'],
])

# Naming Trap
h1('Naming Trap / 命名陷阱')
table([
    ['Context', 'milia_like...'],
    ['mask filename', 'milia_like_cyst (singular / 单数)'],
    ['JSON attributes_order', 'milia_like_cysts (plural / 复数)'],
])

out = 'Task_Requirements_Scoring.docx'
doc.save(out)
print('saved', out)
